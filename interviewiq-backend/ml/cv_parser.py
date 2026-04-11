"""
cv_parser.py — CV/Resume parser with full OCR support

Pipeline:
  PDF (text-based) → pypdf extract → done
  PDF (scanned/image) → pdf2image → pytesseract OCR per page → done
  DOCX → python-docx extract
  Image (PNG/JPG) → pytesseract OCR directly
  TXT → direct decode

Then: extract skills, job titles, experience years, education from text.
"""

import re
import io
import os


def _clean(text: str) -> str:
    # ── Strip PDF operator / metadata garbage ─────────────────────────────
    # Remove PDF content-stream operators that leak into extracted text
    # e.g. "<</Title (Name)>>", "/BM /Normal", "BT ET", "BMC EMC"
    text = re.sub(r'<<[^>]*>>', ' ', text)          # << ... >> PDF dict objects
    text = re.sub(r'/[A-Z][A-Za-z0-9]+ /[A-Z][A-Za-z0-9]+', ' ', text)  # /BM /Normal
    text = re.sub(r'(BT|ET|BMC|EMC|BDC|PDC|BX|EX)', ' ', text)      # PDF operators
    text = re.sub(r'/[A-Za-z][A-Za-z0-9_.:-]{0,30}(?=\s)', ' ', text)    # /PdfName tokens
    text = re.sub(r'\\[0-9]{3}', ' ', text)        # octal escapes \301
    # ── Normalize unicode → ASCII ──────────────────────────────────────────
    text = text.replace('\u00b7', ' ')   # middle dot
    text = text.replace('\u2022', '\n- ')  # bullet → list item
    text = text.replace('\u2019', "'").replace('\u2018', "'")
    text = text.replace('\u201c', '"').replace('\u201d', '"')
    text = text.replace('\u2013', '-').replace('\u2014', '-')
    text = text.replace('\u00e2\u0080\u0099', "'")  # UTF-8 mangled apostrophe
    # ── Remove remaining non-printable chars ──────────────────────────────
    text = re.sub(r'[^\x20-\x7E\n]', ' ', text)
    # ── Normalise whitespace ───────────────────────────────────────────────
    text = re.sub(r'[ \t]{2,}', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    # ── Remove lines that are clearly PDF operators (short, no real words) ─
    lines = []
    for line in text.split('\n'):
        stripped = line.strip()
        # Skip lines that are purely PDF operators / garbage
        if re.match(r'^[/<>\\\[\](){}%0-9. ]{0,40}$', stripped) and not re.search(r'[a-zA-Z]{3,}', stripped):
            continue
        lines.append(line)
    return '\n'.join(lines).strip()


def _preprocess_for_ocr(img):
    """Sharpen + upscale + binarize for better OCR on blurry scans"""
    from PIL import Image, ImageFilter, ImageEnhance

    img = img.convert('L')  # greyscale

    # Upscale if too small — OCR needs ~300 DPI equivalent
    w, h = img.size
    if w < 1400:
        scale = max(2, 1800 // w)
        img = img.resize((w * scale, h * scale), Image.LANCZOS)

    # Sharpen blurry text (apply twice for heavy blur)
    img = img.filter(ImageFilter.SHARPEN)
    img = img.filter(ImageFilter.SHARPEN)

    # Boost contrast
    img = ImageEnhance.Contrast(img).enhance(2.5)

    # Binarize — strict black/white
    img = img.point(lambda p: 255 if p > 140 else 0)

    return img


def _ocr_image(img) -> str:
    import pytesseract
    config = '--psm 3 --oem 3 -c preserve_interword_spaces=1'
    return pytesseract.image_to_string(img, lang='eng', config=config)


def _extract_from_pdf(file_bytes: bytes) -> str:
    text = ''
    direct_text = ''

    # Step 1: try direct text extraction
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or '' for page in reader.pages]
        text = '\n'.join(pages)
        direct_text = text
    except Exception:
        pass

    # Check if text extraction got enough real words
    real_words = re.findall(r'[a-zA-Z]{3,}', text)
    if len(real_words) >= 30:
        return _clean(text)

    # Step 2: OCR fallback — convert each page to image then OCR
    print("  Scanned PDF detected — running OCR...")
    try:
        from pdf2image import convert_from_bytes
        images = convert_from_bytes(file_bytes, dpi=300, fmt='png', thread_count=2)
        pages = []
        for i, img in enumerate(images):
            print(f"    OCR page {i+1}/{len(images)}...")
            processed = _preprocess_for_ocr(img)
            pages.append(_ocr_image(processed))
        text = '\n\n'.join(pages)
        print(f"    Done — {len(text.split())} words extracted")
    except Exception as e:
        print(f"    OCR failed: {e}")
        # Never decode raw PDF bytes as text; that creates massive garbage output
        # and causes false validation failures. Fall back to direct PDF text only.
        fallback_words = re.findall(r'[a-zA-Z]{3,}', direct_text)
        if len(fallback_words) >= 20:
            text = direct_text
        else:
            text = ''

    return _clean(text)


def _extract_from_docx(file_bytes: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return _clean('\n'.join(parts))
    except Exception:
        return file_bytes.decode('utf-8', errors='ignore')


def _extract_from_image(file_bytes: bytes) -> str:
    from PIL import Image
    img = Image.open(io.BytesIO(file_bytes))
    return _clean(_ocr_image(_preprocess_for_ocr(img)))


# ── Skill lists ──────────────────────────────────────────────────────────────

# ── Skill detection strategy ──────────────────────────────────────────────────
#
# Problem: CV skills sections just LIST technology names with no context.
# e.g. "Python, Java, C++, SQL, React, Next.js, Docker, AWS"
# So we CANNOT require surrounding context words — they won't be there.
#
# Solution: Two tiers
# SAFE_SKILLS  — unambiguous tech words, safe to match anywhere in a CV
# RISKY_SKILLS — common English words (go, r, c) that need a nearby tech neighbour
#                "nearby" = within the same line OR within 50 chars in the text
#
# This handles both:
#   - Skills section listings: "C++, Java, Python, SQL"
#   - Prose descriptions: "Built APIs using Django and PostgreSQL"

SAFE_SKILLS = [
    # Programming languages (unambiguous names)
    'python','javascript','typescript','kotlin','golang','rust','scala',
    'matlab','haskell','perl','elixir','fortran','cobol',
    'c++','c#',  # special char languages — handled with custom pattern
    'html','css','xml','json','yaml','markdown',  # markup/data formats
    # Frontend
    'react','react.js','reactjs','angular','svelte','bootstrap','webpack','redux','mobx',
    'gatsby','remix','tailwind','sass','scss','less','jquery','vue',
    'next.js','nextjs','nuxt','nuxtjs','vite','parcel','babel',
    # Backend
    'django','fastapi','flask','laravel','express','nestjs','fastify',
    'spring boot','springboot','asp.net','sinatra','grpc',
    'graphql','rest api','restful',
    # Databases
    'postgresql','postgres','mongodb','mysql','sqlite','redis','sql','nosql',
    'elasticsearch','dynamodb','cassandra','mariadb','firebase',
    'supabase','influxdb','couchdb','neo4j','oracle',
    # Cloud & DevOps
    'aws','azure','gcp','docker','kubernetes','k8s','terraform',
    'ansible','jenkins','nginx','cloudflare','heroku','vercel',
    'netlify','github actions','gitlab ci','circleci','travis',
    'ci/cd','devops','helm','vagrant',
    # Data & AI/ML
    'tensorflow','pytorch','keras','scikit-learn','pandas','numpy',
    'matplotlib','seaborn','jupyter','langchain','openai','huggingface',
    'machine learning','deep learning','computer vision','nlp',
    'natural language processing','data science','data analysis',
    # Mobile
    'react native','flutter','android','ios','swiftui','xcode',
    'android studio','kotlin multiplatform',
    # Tools & platforms
    'git','github','gitlab','bitbucket','jira','confluence','figma',
    'postman','swagger','tableau','power bi','grafana','datadog',
    'splunk','sentry','notion','slack','linux','unix','bash','shell',
    'docker compose','microservices','websocket','oauth','jwt',
    # Testing
    'recharts','chart.js','d3','visx',
    'jest','pytest','selenium','cypress','playwright','junit',
    'mocha','chai','vitest',
]

# Short/ambiguous words — only count if a tech neighbour is nearby in the text
# "nearby" means: within the SAME LINE, or the skill + neighbour within 80 chars
RISKY_SKILLS = {
    'java':  ['python','javascript','typescript','kotlin','maven','gradle','spring',
              'oop','backend','android','programming','language','class','interface',
              'c++','software','developer','engineer','skills','languages','mysql','sql'],
    'c++':   ['python','java','programming','dsa','algorithm','oop','embedded',
              'language','languages','cpp','competitive','systems','software',
              'skills','developer','engineer','javascript','typescript','c'],
    'c#':    ['.net','dotnet','unity','asp','wpf','xamarin','visual studio',
              'microsoft','windows','programming'],
    'c':     ['python','java','c++','programming','embedded','systems',
              'algorithm','language','languages','pointer','memory'],
    'r':     ['python','ggplot','tidyverse','dplyr','rstudio','statistical',
              'data science','analytics','cran','statistics'],
    'go':    ['golang','goroutine','backend','language','programming','grpc',
              'microservice','kubernetes','docker','cloud'],
    'sql':   ['database','mysql','postgres','sqlite','query','relational',
              'mongodb','nosql','table','schema','join','select'],
    'php':   ['laravel','wordpress','composer','symfony','web','backend',
              'apache','mysql','cms','framework'],
    'ruby':  ['gem','bundler','sinatra','web','backend','heroku',
              'activerecord','rspec','rake','ruby on rails'],
    # 'rails' removed — too ambiguous, use 'ruby on rails' instead
    'sql':   ['database','mysql','postgres','postgresql','sqlite','query',
              'relational','nosql','mongodb','table','schema','select','join',
              'orm','migration','backend','data','python','javascript','java',
              'react','django','flask','web','development','programming'],
    'swift': ['ios','xcode','apple','uikit','swiftui','objective','cocoa',
              'iphone','ipad','mobile','app store'],
    'node':  ['javascript','nodejs','npm','express','backend','server',
              'typescript','react','fullstack','api'],
    'vue':   ['javascript','frontend','nuxt','vuex','component','spa',
              'web','framework','typescript'],
}

SOFT_SKILLS = [
    'leadership','communication','teamwork','problem solving','agile','scrum',
    'project management','mentoring','critical thinking','collaboration',
]

SKILL_CANONICAL_MAP = {
    'react.js': 'react',
    'reactjs': 'react',
    'nextjs': 'next.js',
    'nuxtjs': 'nuxt',
    'springboot': 'spring boot',
    'postgres': 'postgresql',
    'k8s': 'kubernetes',
    'restful': 'rest api',
}

TITLE_PATTERNS = [
    r'(senior|junior|lead|principal|staff|associate|mid|entry.level)?\s*'
    r'(software|frontend|front.end|backend|back.end|full.?stack|mobile|ios|android|'
    r'devops|sre|platform|ml|ai|data|cloud|security|qa|test|systems?)?\s*'
    r'(engineer|developer|architect|analyst|scientist|manager|director|lead|intern|consultant|specialist)',
    r'(product|project|engineering|technical|chief)\s*(manager|director|lead|officer|head)',
    r'(cto|ceo|coo|vp|head of \w+)',
    r'(ui|ux|product|graphic)\s*(designer|researcher|manager|lead)',
    r'(data|business|marketing|systems)\s*(analyst|scientist|engineer)',
]

EDU_KEYWORDS = [
    'bachelor','master','phd','ph.d','bsc','msc','mba','be','btech','mtech',
    'university','college','institute','computer science','software engineering',
    'information technology','electrical engineering','mathematics','graduated','gpa','cgpa',
]


def _extract_skills(text: str) -> list:
    tl = text.lower()
    lines = tl.split('\n')
    found = []

    # ── Tier 1: Safe unambiguous skills — smart pattern per skill ───────────
    for skill in SAFE_SKILLS:
        esc = re.escape(skill)
        # Skills ending in special chars (c++, c#, .net) can't use \b at end
        # Skills starting with . (.net, next.js) can't use \b at start
        if skill[-1] in ('+', '#', '.') or skill[0] == '.':
            pattern = r'(?<![a-zA-Z0-9])' + esc + r'(?![a-zA-Z0-9])'
        else:
            pattern = r'\b' + esc + r'\b'
        if re.search(pattern, tl, re.IGNORECASE):
            found.append(skill)

    # ── Tier 2: Risky short/common words — need a tech neighbour ─────────────
    for skill, neighbours in RISKY_SKILLS.items():
        pattern = r'\b' + re.escape(skill) + r'\b'
        if not re.search(pattern, tl):
            continue

        skill_found = False

        # Check 1: same line as a tech neighbour
        for line in lines:
            if re.search(pattern, line):
                if any(n in line for n in neighbours):
                    skill_found = True
                    break

        # Check 2: within 120 chars of a tech neighbour anywhere in text
        if not skill_found:
            for m in re.finditer(pattern, tl):
                start = max(0, m.start() - 120)
                end   = min(len(tl), m.end() + 120)
                window = tl[start:end]
                if any(n in window for n in neighbours):
                    skill_found = True
                    break

        if skill_found:
            found.append(skill)

    # ── Tier 3: Soft skills — only in longer documents ────────────────────────
    if len(tl.split()) > 60:
        for skill in SOFT_SKILLS:
            if re.search(r'\b' + re.escape(skill) + r'\b', tl):
                found.append(skill)

    # Canonicalize aliases and remove noisy/non-ATS generic words
    canonical = []
    seen = set()
    noisy = {'less', 'expressive'}
    for skill in found:
        s = SKILL_CANONICAL_MAP.get(skill, skill)
        if s in noisy:
            continue
        if s not in seen:
            seen.add(s)
            canonical.append(s)
    return canonical


def _extract_titles(text):
    tl = text.lower()
    titles = []
    for pat in TITLE_PATTERNS:
        for m in re.finditer(pat, tl, re.IGNORECASE):
            t = re.sub(r'\s+', ' ', m.group(0)).strip()
            if 4 < len(t) < 60:
                titles.append(t.title())
    return list(dict.fromkeys(titles))[:6]


def _extract_exp_years(text):
    tl = text.lower()
    # Explicit mention: "5 years of experience"
    m = re.search(r'(\d+)\+?\s*years?\s*(of\s*)?(experience|work|exp\.?)', tl)
    if m:
        return min(int(m.group(1)), 35)
    # Estimate from work-experience section only (avoids education/project year inflation)
    lines = text.split('\n')
    section = []
    capture = False
    for ln in lines:
        l = ln.strip().lower()
        if re.search(r'\b(work\s+experience|experience|employment|work history)\b', l):
            capture = True
            continue
        if capture and re.search(r'\b(education|skills|projects|certifications|summary|objective|profile)\b', l):
            break
        if capture:
            section.append(ln)

    candidate_text = '\n'.join(section) if section else text
    if not re.search(r'\b(company|intern|engineer|developer|worked|role|position|employer)\b', candidate_text.lower()):
        return 0

    date_ranges = re.findall(
        r'\b(19[8-9]\d|20[0-2]\d)\b\s*[-–]\s*(present|\b(19[8-9]\d|20[0-2]\d)\b)',
        candidate_text.lower()
    )
    if date_ranges:
        current_year = 2026
        spans = []
        for start, end_text, _ in date_ranges:
            s = int(start)
            e = current_year if end_text == 'present' else int(end_text)
            if s <= e:
                spans.append(e - s)
        if spans:
            return max(0, min(max(spans), 35))
    return 0


def _extract_education(text):
    edu = []
    for line in text.split('\n'):
        if any(k in line.lower() for k in EDU_KEYWORDS) and 8 < len(line.strip()) < 250:
            edu.append(line.strip())
    return list(dict.fromkeys(edu))[:4]


def _extract_name(text: str) -> str:
    """
    Extract candidate name. Handles:
    - Clean name on first line: "Miral Sajid"
    - Name with contact on same line: "Bakhtawar Baig\n03094..."
    - After PDF metadata cleanup still has issues: skip lines with / < > \\ [ ]
    """
    SKIP_WORDS = {
        'resume','cv','curriculum','vitae','profile','summary','objective',
        'experience','education','skills','contact','references','linkedin',
        'github','portfolio','email','phone','address','software','engineer',
        'developer','manager','analyst','designer','consultant','senior','junior',
        'bachelor','master','university','college','government','national',
    }

    lines_raw = [l.strip() for l in text.split('\n') if l.strip()]
    # Remove lines with PDF garbage characters
    lines = [l for l in lines_raw if not re.search(r'[/<>\\\[\]{}]', l)]

    for line in lines[:15]:
        # Skip if contains digits, @, |, URLs, or is too long
        if re.search(r'[@|·•\d]', line):
            continue
        if re.search(r'https?://', line, re.I):
            continue
        if len(line) > 60:
            continue
        words = line.split()
        if not (2 <= len(words) <= 4):
            continue
        # Every word: starts capital, only letters, 2-25 chars
        if not all(w[0].isupper() and w.isalpha() and 2 <= len(w) <= 25 for w in words):
            continue
        line_lower = line.lower()
        if any(skip in line_lower for skip in SKIP_WORDS):
            continue
        return line

    # Fallback: name may be on a line with phone/email — extract just the name part
    for line in lines[:8]:
        m = re.match(r'^([A-Z][a-z]+(?: [A-Z][a-z]+){1,3})', line)
        if m:
            candidate = m.group(1)
            if not any(skip in candidate.lower() for skip in SKIP_WORDS):
                return candidate

    return ''


def _extract_email(text):
    m = re.search(r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}', text)
    return m.group(0) if m else ''


def _extract_phone(text):
    m = re.search(r'(\+?\d[\d\s\-().]{8,15}\d)', text)
    return m.group(0).strip() if m else ''


def _extract_summary(text):
    m = re.search(
        r'(?:summary|objective|profile|about me)[:\s]*\n?([\s\S]{50,400}?)(?:\n[A-Z]{3,}|\n\n)',
        text, re.IGNORECASE
    )
    return m.group(1).strip()[:400] if m else ''


# ── ATS Resume Validation ─────────────────────────────────────────────────────

# Signals that strongly indicate this is a CV/resume
CV_SIGNALS = [
    'experience', 'education', 'skills', 'objective', 'summary', 'profile',
    'employment', 'work history', 'career', 'references', 'achievements',
    'certifications', 'qualifications', 'responsibilities', 'projects',
    'internship', 'bachelor', 'master', 'university', 'college', 'degree',
    'gpa', 'cgpa', 'resume', 'curriculum vitae', 'cv', 'worked at',
    'position', 'employer', 'job title', 'linkedin', 'github.com',
]

# Signals that indicate this is NOT a CV (assignment, report, essay etc.)
NON_CV_SIGNALS = [
    'question', 'answer', 'marks', 'assignment', 'submitted by', 'due date',
    'instructor', 'professor', 'course code', 'lecture', 'homework',
    'problem statement', 'solution', 'given that', 'calculate', 'prove that',
    'hypothesis', 'methodology', 'bibliography', 'references:', 'chapter',
    'abstract:', 'introduction:', 'conclusion:', 'table of contents',
    'figure 1', 'figure 2', 'equation', 'theorem', 'lemma', 'appendix',
    'word count:', 'submitted to', 'page number',
    'mcq', 'multiple choice', 'semester', 'roll no', 'student id',
]

ATS_SECTIONS = ['experience', 'education', 'skills', 'work', 'employment',
                'projects', 'certifications', 'summary', 'objective', 'profile']

# Regex variants for common resume headings and wording differences
ATS_SECTION_PATTERNS = [
    r'\b(work\s+)?experience\b',
    r'\bemployment(\s+history)?\b',
    r'\beducation(al)?\b',
    r'\bskills?\b',
    r'\btechnical\s+skills?\b',
    r'\bprojects?\b',
    r'\bcertifications?\b',
    r'\bsummary\b',
    r'\bprofessional\s+summary\b',
    r'\bobjective\b',
    r'\bprofile\b',
]


def validate_cv(raw_text: str) -> dict:
    """
    Validate whether uploaded file is a resume/CV (not an assignment/report).

    Strategy:
    - Count CV-positive signals (experience, education, skills sections etc.)
    - Count CV-negative signals (assignment, question, marks, due date etc.)
    - Make decision based on ratio, NOT word count (OCR inflates word count)
    - Only hard-reject if clearly not a resume
    """
    tl = raw_text.lower()
    # Use unique words sample for signal detection (avoids OCR repetition inflation)
    unique_words = set(tl.split())
    warnings = []

    # ── Hard reject: almost no real text at all ───────────────────────────
    real_words = [w for w in unique_words if len(w) >= 3 and w.isalpha()]
    if len(real_words) < 30:
        return {'is_valid': False, 'score': 0,
                'reason': 'Could not extract enough text from this document. '
                           'If it is a scanned PDF, ensure the scan quality is good. '
                           'Try saving it as a text-based PDF or DOCX instead.',
                'warnings': []}

    # ── Check non-CV signals (assignment/academic doc) ────────────────────
    # Only count signals that appear as distinct words (not substrings)
    non_cv_hits = []
    for sig in NON_CV_SIGNALS:
        # Multi-word signals: just check if present in text
        if ' ' in sig:
            if sig in tl:
                non_cv_hits.append(sig)
        else:
            if re.search(r'\b' + re.escape(sig) + r'\b', tl):
                non_cv_hits.append(sig)

    # ── Count CV signals ──────────────────────────────────────────────────
    cv_hits = sum(1 for sig in CV_SIGNALS if sig in tl)
    section_hits = sum(1 for pat in ATS_SECTION_PATTERNS if re.search(pat, tl))

    # Resume structure fallback signals (for strict/edge OCR cases)
    has_email = bool(re.search(r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}', raw_text))
    has_phone = bool(re.search(r'(\+?\d[\d\s\-().]{8,15}\d)', raw_text))
    has_profile_link = ('linkedin' in tl) or ('github' in tl)
    has_date_ranges = bool(
        re.search(r'\b(19[8-9]\d|20[0-2]\d)\b\s*[-–]\s*(present|\b(19[8-9]\d|20[0-2]\d)\b)', tl)
    )
    has_education_terms = any(k in tl for k in EDU_KEYWORDS)
    skills_preview = _extract_skills(raw_text)
    has_skill_density = len(skills_preview) >= 4

    fallback_resume_score = sum([
        has_email,
        has_phone,
        has_profile_link,
        has_date_ranges,
        has_education_terms,
        has_skill_density,
    ])

    # Resume structure fallback signals (for strict/edge OCR cases)
    # ...existing fallback setup above...

    # Strong "resume shape" checks
    heading_hits = len(re.findall(
        r'(?m)^\s*(work experience|experience|education|skills|projects|certifications|summary|objective|profile)\s*$',
        tl
    ))
    has_resume_contacts = has_email or has_phone or has_profile_link

    # ── Decision logic ────────────────────────────────────────────────────
    # A resume typically has: many CV signals, few non-CV signals
    # An assignment typically has: many non-CV signals, few CV signals

    # Strong rejection: lots of academic signals AND very few resume signals
    if len(non_cv_hits) >= 4 and cv_hits <= 2:
        return {'is_valid': False, 'score': 0,
                'reason': f'This looks like an academic document (assignment/report), not a resume. '
                           f'Please upload your CV only. '
                           f'Detected: {", ".join(non_cv_hits[:4])}',
                'warnings': []}

    # Medium rejection: clearly academic with no resume structure at all
    if len(non_cv_hits) >= 5 and section_hits == 0:
        return {'is_valid': False, 'score': 0,
                'reason': 'This does not appear to be a resume. No resume sections found. ' 
                           'Please upload a CV with Work Experience, Education, and Skills.',
                'warnings': []}

    # Reject docs that look like assignments/notes even if some tech words exist
    if len(non_cv_hits) >= 4 and not has_resume_contacts and heading_hits < 2:
        return {'is_valid': False, 'score': 0,
                'reason': 'This appears to be an assignment/report, not a CV. '
                          'Please upload a resume with contact details and CV sections.',
                'warnings': []}

    # Needs SOME resume content to be valid
    confidence = min(100, cv_hits * 8 + section_hits * 12)
    if cv_hits < 2 and section_hits < 1 and fallback_resume_score < 3:
        return {'is_valid': False, 'score': confidence,
                'reason': 'This does not appear to be a resume. '
                           'Please upload a CV that includes Work Experience, Education, and Skills sections.',
                'warnings': []}

    # ── ATS warnings (non-blocking) ───────────────────────────────────────
    if not re.search(r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}', raw_text):
        warnings.append('No email found — add your email for ATS systems')
    if not re.search(r'\bskills\b', tl):
        warnings.append('No Skills section detected — add one for better ATS parsing')

    return {'is_valid': True, 'score': confidence, 'reason': '', 'warnings': warnings}


# ── Main function ─────────────────────────────────────────────────────────────

def parse_cv(file_bytes: bytes, filename: str) -> dict:
    fn  = (filename or '').lower()
    ext = os.path.splitext(fn)[1]

    print(f"Parsing: {filename} ({len(file_bytes)//1024} KB)")

    if ext == '.pdf':
        raw = _extract_from_pdf(file_bytes)
    elif ext in ('.docx', '.doc'):
        raw = _extract_from_docx(file_bytes)
    elif ext in ('.png', '.jpg', '.jpeg', '.webp', '.bmp', '.tiff'):
        raw = _extract_from_image(file_bytes)
    else:
        raw = _clean(file_bytes.decode('utf-8', errors='ignore'))

    print(f"  Text: {len(raw.split())} words")

    # ── Validate before extracting ────────────────────────────────────────
    validation = validate_cv(raw)
    if not validation['is_valid']:
        return {
            'is_valid':   False,
            'error':      validation['reason'],
            'score':      validation['score'],
            'raw_text':   raw[:500],
            'word_count': len(raw.split()),
        }

    skills    = _extract_skills(raw)
    titles    = _extract_titles(raw)
    exp_years = _extract_exp_years(raw)
    education = _extract_education(raw)
    name      = _extract_name(raw)
    summary   = _extract_summary(raw)

    parts = []
    if titles:    parts.append(f"Roles: {', '.join(titles[:3])}")
    if exp_years: parts.append(f"Experience: ~{exp_years} years")
    if skills:    parts.append(f"Skills: {', '.join(skills[:12])}")
    if education: parts.append(f"Education: {education[0][:100]}")
    if summary:   parts.append(f"Profile: {summary[:200]}")

    print(f"  Result: {len(skills)} skills, {len(titles)} titles, {exp_years}yr exp")

    # Count only real alphabetic words (not OCR noise) for display
    clean_word_count = len([w for w in raw.split() if len(w) >= 2 and any(c.isalpha() for c in w)])

    return {
        'is_valid':         True,
        'warnings':         validation['warnings'],
        'name':             name,
        'email':            _extract_email(raw),
        'phone':            _extract_phone(raw),
        'skills':           skills[:15],
        'job_titles':       titles,
        'experience_years': exp_years,
        'education':        education,
        'professional_summary': summary,
        'cv_summary':       '. '.join(parts),
        'raw_text':         raw[:4000],
        'word_count':       clean_word_count,
    }
